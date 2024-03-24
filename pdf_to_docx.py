import PyPDF2
from docx import Document
from PIL import Image
import os

def pdf_to_word(pdf_file, word_file):
    # PDF 파일 열기
    with open(pdf_file, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        # Word 문서 생성
        doc = Document()

        # 이미지를 저장할 디렉토리 생성
        img_dir = 'images'
        if not os.path.exists(img_dir):
            os.makedirs(img_dir)

        # PDF 페이지를 반복하며 텍스트와 이미지를 추출하여 Word 문서에 추가
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            doc.add_paragraph(text)
            
            # 이미지 추출 및 삽입
            if '/XObject' in page['/Resources']:
                xObject = page['/Resources']['/XObject'].get_object()
                for obj in xObject:
                    if xObject[obj]['/Subtype'] == '/Image':
                        size = (xObject[obj]['/Width'], xObject[obj]['/Height'])
                        data = xObject[obj].get_data()
                        if xObject[obj]['/ColorSpace'] == '/DeviceRGB':
                            mode = 'RGB'
                        else:
                            mode = 'P'
                        if xObject[obj]['/Filter'] == '/FlateDecode':
                            img = Image.frombytes(mode, size, data)
                            img_path = os.path.join(img_dir, f'page_{page_num}_image_{obj}.png')
                            img.save(img_path)
                            doc.add_picture(img_path)
                        elif xObject[obj]['/Filter'] == '/DCTDecode':
                            img_path = os.path.join(img_dir, f'page_{page_num}_image_{obj}.jpg')
                            with open(img_path, 'wb') as img_file:
                                img_file.write(data)
                            doc.add_picture(img_path)
                        elif xObject[obj]['/Filter'] == '/JPXDecode':
                            img_path = os.path.join(img_dir, f'page_{page_num}_image_{obj}.jp2')
                            with open(img_path, 'wb') as img_file:
                                img_file.write(data)
                            doc.add_picture(img_path)
        
        # Word 파일로 저장
        doc.save(word_file)
        
        print("PDF를 Word 파일로 변환 완료!")

# PDF 파일과 Word 파일의 경로 설정
pdf_file = 'example.pdf'
word_file = 'example.docx'

# PDF를 Word로 변환
pdf_to_word(pdf_file, word_file)
